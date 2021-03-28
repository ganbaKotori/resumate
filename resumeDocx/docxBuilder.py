from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH 


userInformation = {
        "education": [{
            "city": "Long Beach",
            "endDate": "2020-05-31T03:09:23.488Z",
            "schoolName": "California State University, Long Beach",
            "startDate": "2017-08-27T03:09:23.488Z",
            "state": "CA",
            "degree": "bachelors",
            "major": "computer science",
        }],
        "email": "ajramirez1095@gmail.com",
        "experience": [{
            "city": "El Segundo",
            "companyName": "LinQuest",
            "description": "Developed SharePoint webparts using JavaScript",
            "endDate": "2019-08-27T03:09:23.488Z",
            "position": "Web Developer, Intern",
            "startDate": "2019-06-21T03:09:23.488Z",
            "state": "CA"
        }],
        "firstName": "Alexander",
        "lastName": "Ramirez",
        "phoneNumber": "5625837458",
        "skills": ["JavaScript", "HTML/CSS", "jQuery"],
        #"url": "https://www.indeed.com/viewjob?cmp=Tapia-Brothers-Company&t=Software+Engineer&jk=caa43b6de8a2b3b4&sjdu=QwrRXKrqZ3CNX5W-O9jEvRFd8FQI4DEv5V74lSpSnHYuqC9T3Ssm3_JS_P7tvu6NVVYgJPClrjl-TU2ErV4ICA&tk=1f1pvqo72oc90800&adid=358016341&ad=-6NYlbfkN0DSJVOgFNjTH3NHsz1y5W7iEiQ8XFPauIDjgJYsd-OdNGsRX_DeFFbkppNd1eMCXipH7veJakjemwe992tJiEu_aZIfWMKYlkH5qgcuSYrIBDrPtfDoQj22TNBQtlQH2tT3xV17d1aht6MTX29Jq-LnTaNNL5NRwCsv_NBrVxaAoL5nh4Bg6T95OBvmm-Y2OuQr_ZovCwZHC8P9MxtA9t8_rGE9ceAd2N5CiegfvGtAStprPXDEZJvlUXp4-UjyUBKoaPyoClUe-XhWC71Z0XBcagzY6nPNZlmJr46P7Jfqg9VLROWEdfB60mlcuV2Y3PSpMJ9z82PLAe-tJS2YdZwt&pub=4a1b367933fd867b19b072952f68dceb&vjs=3"
    }


#----------------------------------------------------------------------#
document = Document()
sections = document.sections

#size margins
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

#main header, 1 is equal to center alignment
mainHeader = document.add_heading(userInformation["firstName"] + " " + userInformation["lastName"])
mainHeader.alignment = 1

#contact info of the user 
contactInfo = document.add_paragraph('Phone: ' + userInformation["phoneNumber"])
contactInfo.add_run('\t'.expandtabs(8))
contactInfo.add_run('Email: ' + userInformation["email"])

#experience of the user 
document.add_heading('Experience', level=1)
for i in userInformation["experience"]:
    expInfo= document.add_paragraph (i["position"])
    expInfo.add_run("                                ")
    expInfo.add_run(i["companyName"] + ", " + i["city"])
    expInfo.add_run("                         ")
    #added info
    expInfo.add_run('June to August 2019')
    document.add_paragraph('Developed SharePoint webparts using HTML, CSS, JavaScript, AJAX, and jQuery', style='List Bullet')
    document.add_paragraph('Practiced Agile Software Development by attending Scrum Meetings and Sprint Reviews', style='List Bullet')
    document.add_paragraph('Replaced XSLT webparts on team’s Kanban page with FusionCharts webparts using JavaScript and jQuery', style='List Bullet')
    document.add_paragraph('Improved User Experience and Loading Time of Kanban page with replacement of webparts', style='List Bullet')
    
    #document.add_paragraph(i["startDate"] + " to " + i["endDate"])

#education of the user 
document.add_heading('Education', level=1)
for i in userInformation["education"]:
    eduInfo= document.add_paragraph(i["schoolName"])
    eduInfo.add_run("                                           ")
    #added info
    eduInfo.add_run('August 2017 to May 2020')
    #added major from data
    document.add_paragraph(i["degree"] + ": " + i["major"])


#adding in selected projects, couldn't fine info so added in manually
document.add_heading('Selected Projects', level=1)
projectExp = document.add_paragraph('Bluefin: Los Angeles')
projectExp.add_run = ("                                ")
projectExp.add_run = ('January 2021 to Current')
document.add_paragraph('Created a website that displays location-based data such as median income, crime spots, schools, and diversity within Los Angeles', style='List Bullet')
document.add_paragraph('Pulled and filtered data using jQuery from the City of Los Angeles’s GeoHub website', style='List Bullet')
document.add_paragraph('Utilized Leaflet to pinpoint each datapoint with markers and polygons', style='List Bullet')
document.add_paragraph('Implemented 3rd party Leaflet libraries such as an address search bar and a Find My Location feature', style='List Bullet')

#added in user skills
document.add_heading('Skills', level=1)
document.add_paragraph(userInformation["skills"])


document.add_page_break()

document.save('resume.docx')