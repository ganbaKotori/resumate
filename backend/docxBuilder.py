from docx import Document
from docx.shared import Inches, Cm
import datetime
#from docx.enum.text import WD_ALIGN_PARAGRAPH
# -*- coding: utf-8 -*-
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
date_format = '%Y-%m-%d'


userInformation = {
        "education": [{
            "city": "Long Beach",
            "endDate": "2020-05-31T03:09:23.488Z",
            "schoolName": "California State University, Long Beach",
            "startDate": "2017-08-27T03:09:23.488Z",
            "state": "CA",
            "degree": "bachelors",
            "major": "computer science",
        }],
        "email": "ajramirez1095@gmail.com",
        "experience": [{
            "city": "El Segundo",
            "companyName": "LinQuest",
            "description": "Developed SharePoint webparts using JavaScript",
            "endDate": "2019-08-27T03:09:23.488Z",
            "position": "Web Developer, Intern",
            "startDate": "2019-06-21T03:09:23.488Z",
            "state": "CA"
        }],
        "firstName": "Alexander",
        "lastName": "Ramirez",
        "phoneNumber": "5625837458",
        "skills": ["JavaScript", "HTML/CSS", "jQuery"],
        
    }


#----------------------------------------------------------------------#
def buildResume(userInformation):
    document = Document()
    sections = document.sections

    #size margins
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    #main header, 1 is equal to center alignment
    mainHeader = document.add_heading(userInformation["firstName"] + " " + userInformation["lastName"])
    mainHeader.alignment = 1

    #contact info of the user 
    contactInfo = document.add_paragraph('Phone: ' + userInformation["phoneNumber"])
    contactInfo.add_run('\t'.expandtabs(8))
    contactInfo.add_run('Email: ' + userInformation["email"])

    #experience of the user 
    document.add_heading('Experience', level=1)
    for i in userInformation["experience"]:
        try:
            startDate = datetime.datetime.strptime(i["startDate"][0:10], date_format)
            endDate = datetime.datetime.strptime(i["endDate"][0:10], date_format)
        except ValueError:
            print("Incorrect data format, should be YYYY-MM-DD")

        expInfo= document.add_paragraph (i["position"])
        expInfo.add_run("                                ")
        expInfo.add_run(i["companyName"] + ", " + i["city"])
        expInfo.add_run("                         ")
        #added info
        #expInfo.add_run(startDate.strftime("%B") + ' ' + startDate.strftime("%d") + ',' + startDate.strftime("%Y") + ' to ' + startDate.strftime("%B") + ' ' + startDate.strftime("%d") + ',' + startDate.strftime("%Y"))
        expInfo.add_run(startDate.strftime("%B") + ' '  + startDate.strftime("%Y") + ' to ' + endDate.strftime("%B") + ' ' + endDate.strftime("%Y"))
        document.add_paragraph(i["description"], style='List Bullet')
        #document.add_paragraph('Practiced Agile Software Development by attending Scrum Meetings and Sprint Reviews', style='List Bullet')
        #document.add_paragraph('Coded', style='List Bullet')
        #document.add_paragraph('Improved User Experience and Loading Time of Kanban page with replacement of webparts', style='List Bullet')
        
        #document.add_paragraph(i["startDate"] + " to " + i["endDate"])

    #education of the user 
    document.add_heading('Education', level=1)
    for i in userInformation["education"]:
        try:
            startDate2 = datetime.datetime.strptime(i["startDate"][0:10], date_format)
            endDate2 = datetime.datetime.strptime(i["endDate"][0:10], date_format)
        except ValueError:
            print("Incorrect data format, should be YYYY-MM-DD")
        eduInfo= document.add_paragraph(i["schoolName"])
        eduInfo.add_run("                                           ")
        #added info
        eduInfo.add_run(startDate2.strftime("%B") + ' '  + startDate2.strftime("%Y") + ' to ' + endDate2.strftime("%B") + ' ' + endDate2.strftime("%Y"))
        #added major from data
        document.add_paragraph(i["degree"] + " in " + i["major"], style='List Bullet')


    #adding in selected projects, couldn't fine info so added in manually
    #document.add_heading('Selected Projects', level=1)
    #projectExp = document.add_paragraph('Bluefin: Los Angeles')
    #projectExp.add_run = ("                                ")
    #projectExp.add_run = ('January 2021 to Current')
    #document.add_paragraph('Created a website that displays location-based data such as median income, crime spots, schools, and diversity within Los Angeles', style='List Bullet')
    #document.add_paragraph('Pulled and filtered data using jQuery from the City of Los Angeless GeoHub website', style='List Bullet')
    #document.add_paragraph('Utilized Leaflet to pinpoint each datapoint with markers and polygons', style='List Bullet')
    #document.add_paragraph('Implemented 3rd party Leaflet libraries such as an address search bar and a Find My Location feature', style='List Bullet')

    #added in user skills
    document.add_heading('Skills', level=1)
    skills = ", ".join(userInformation["skills"])
    document.add_paragraph(skills)


    document.add_page_break()

    document.save('./resumes/resume.docx')

#buildResume(userInformation)